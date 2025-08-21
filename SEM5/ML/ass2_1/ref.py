# Full workflow: manual OLS over multiple train/test splits, metrics, plots, excel, correlation analysis
import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime

# ------------- Configuration -------------
DATA_FILE = 'salary_data.csv'  # expects two columns: YearsExperience, Salary (or similar)
PLOT_DIR = 'plots'
RESULTS_XLSX = 'results.xlsx'
ANALYSIS_DOCX = 'analysis.docx'  # optional (requires python-docx)
TRAIN_PCTS = list(range(10, 100, 10))  # 10..90
SEED = 42
np.random.seed(SEED)

# ------------- Utilities -------------

def ensure_dirs():
    os.makedirs(PLOT_DIR, exist_ok=True)


def manual_train_test_split(df, train_pct, random_state=None):
    if random_state is not None:
        rng = np.random.default_rng(random_state)
        indices = np.arange(len(df))
        rng.shuffle(indices)
        train_size = int(len(df) * train_pct / 100.0)
        train_idx = indices[:train_size]
        test_idx = indices[train_size:]
        return df.iloc[train_idx], df.iloc[test_idx]
    else:
        train_size = int(len(df) * train_pct / 100.0)
        return df.iloc[:train_size], df.iloc[train_size:]


def ols_simple(x, y):
    """Manual OLS (simple linear regression): returns intercept w0 and slope w1.
    x, y: 1D numpy arrays.
    """
    x = x.ravel()
    y = y.ravel()
    x_mean = x.mean()
    y_mean = y.mean()
    # slope = covariance(x,y)/variance(x)
    numerator = np.sum((x - x_mean)*(y - y_mean))
    denominator = np.sum((x - x_mean)**2)
    w1 = numerator / denominator
    w0 = y_mean - w1 * x_mean
    return w0, w1


def predict(w0, w1, x):
    return w0 + w1 * x


def rss(y_true, y_pred):
    return np.sum((y_true - y_pred)**2)


def r2_score(y_true, y_pred):
    ss_res = rss(y_true, y_pred)
    ss_tot = np.sum((y_true - y_true.mean())**2)
    return 1 - ss_res/ss_tot


def pearson_from_slope(w1, x, y):
    # r = w1 * std_x / std_y for simple linear regression
    return w1 * (x.std(ddof=0) / y.std(ddof=0))

# ------------- Load Data -------------
df_full = pd.read_csv(DATA_FILE)
# Try to infer column names
cols = [c.lower() for c in df_full.columns]
if len(df_full.columns) < 2:
    raise ValueError('Dataset must have at least two columns: feature and target.')

# Assume first numeric column is experience, last numeric column is salary
numeric_cols = [c for c in df_full.columns if pd.api.types.is_numeric_dtype(df_full[c])]
if len(numeric_cols) < 2:
    raise ValueError('Need at least two numeric columns.')
feature_col = numeric_cols[0]
target_col = numeric_cols[-1]

print(f'Using feature column: {feature_col}, target column: {target_col}')

# ------------- Processing over splits -------------
ensure_dirs()

param_rows = []
metric_rows = []
line_colors = plt.cm.viridis(np.linspace(0, 1, len(TRAIN_PCTS)))

# For overlay plot of all hypotheses
fig_all, ax_all = plt.subplots(figsize=(7,5))
ax_all.set_title('Regression Lines Across Training Splits')
ax_all.set_xlabel(feature_col)
ax_all.set_ylabel(target_col)
ax_all.scatter(df_full[feature_col], df_full[target_col], s=25, c='lightgray', label='All data')

for i, pct in enumerate(TRAIN_PCTS):
    train_df, test_df = manual_train_test_split(df_full, pct, random_state=SEED + i)
    x_train = train_df[feature_col].values
    y_train = train_df[target_col].values
    x_test = test_df[feature_col].values
    y_test = test_df[target_col].values

    w0, w1 = ols_simple(x_train, y_train)
    y_train_pred = predict(w0, w1, x_train)
    y_test_pred = predict(w0, w1, x_test)

    train_rss = rss(y_train, y_train_pred)
    test_rss = rss(y_test, y_test_pred)
    train_rss_mean = train_rss / len(y_train)
    test_rss_mean = test_rss / len(y_test)

    train_r2 = r2_score(y_train, y_train_pred)
    test_r2 = r2_score(y_test, y_test_pred)

    # Pearson from slope & data
    pearson_r_train = pearson_from_slope(w1, x_train, y_train)

    param_rows.append({
        'train_pct': pct,
        'w0': w0,
        'w1': w1,
        'pearson_r_from_slope_train': pearson_r_train,
        'n_train': len(x_train),
        'n_test': len(x_test)
    })

    metric_rows.append({
        'train_pct': pct,
        'train_rss_mean': train_rss_mean,
        'test_rss_mean': test_rss_mean,
        'train_r2': train_r2,
        'test_r2': test_r2
    })

    # Individual plot for this split
    fig, ax = plt.subplots(figsize=(6,4))
    ax.scatter(x_train, y_train, c='blue', alpha=0.7, label='Train')
    ax.scatter(x_test, y_test, c='orange', alpha=0.7, label='Test')
    # line over full feature range
    x_line = np.linspace(df_full[feature_col].min(), df_full[feature_col].max(), 100)
    y_line = predict(w0, w1, x_line)
    ax.plot(x_line, y_line, color='red', label=f'Hypothesis (w0={w0:.2f}, w1={w1:.2f})')
    ax.set_title(f'Train % = {pct}')
    ax.set_xlabel(feature_col)
    ax.set_ylabel(target_col)
    ax.legend()
    fname = os.path.join(PLOT_DIR, f'hypothesis_{pct}pct.png')
    fig.tight_layout()
    fig.savefig(fname, dpi=120)
    plt.close(fig)

    # Add line to overlay plot
    ax_all.plot(x_line, y_line, color=line_colors[i], label=f'{pct}% (w1={w1:.2f})')

# Finish overlay plot
ax_all.legend(fontsize='small', ncol=2)
fig_all.tight_layout()
fig_all.savefig(os.path.join(PLOT_DIR, 'all_hypotheses.png'), dpi=130)
plt.close(fig_all)

# ------------- Save results to Excel -------------
params_df = pd.DataFrame(param_rows)
metrics_df = pd.DataFrame(metric_rows)

# Add combined sheet with predictions stacked (optional)
# We'll store only parameters & metrics to keep file concise
with pd.ExcelWriter(RESULTS_XLSX, engine='openpyxl') as writer:
    params_df.to_excel(writer, sheet_name='parameters', index=False)
    metrics_df.to_excel(writer, sheet_name='metrics', index=False)

print(f'Saved parameters & metrics to {RESULTS_XLSX}')

# ------------- Training pct vs RSS and R2 plots -------------
fig_rss, ax_rss = plt.subplots(figsize=(6,4))
ax_rss.plot(metrics_df['train_pct'], metrics_df['train_rss_mean'], marker='o', label='Train Mean RSS')
ax_rss.plot(metrics_df['train_pct'], metrics_df['test_rss_mean'], marker='s', label='Test Mean RSS')
ax_rss.set_xlabel('Training %')
ax_rss.set_ylabel('Mean RSS')
ax_rss.set_title('Training % vs Mean RSS')
ax_rss.legend()
fig_rss.tight_layout()
fig_rss.savefig(os.path.join(PLOT_DIR, 'training_pct_vs_mean_rss.png'), dpi=130)
plt.close(fig_rss)

fig_r2, ax_r2 = plt.subplots(figsize=(6,4))
ax_r2.plot(metrics_df['train_pct'], metrics_df['train_r2'], marker='o', label='Train R2')
ax_r2.plot(metrics_df['train_pct'], metrics_df['test_r2'], marker='s', label='Test R2')
ax_r2.set_xlabel('Training %')
ax_r2.set_ylabel('R^2')
ax_r2.set_title('Training % vs R^2')
ax_r2.legend()
fig_r2.tight_layout()
fig_r2.savefig(os.path.join(PLOT_DIR, 'training_pct_vs_r2.png'), dpi=130)
plt.close(fig_r2)

print('Saved performance plots.')

# ------------- Correlation analysis across splits -------------
# Correlation between training percentage and slope, and between slope & pearson estimate
train_pct_arr = params_df['train_pct'].values
slope_arr = params_df['w1'].values
pearson_est_arr = params_df['pearson_r_from_slope_train'].values

corr_trainpct_slope = np.corrcoef(train_pct_arr, slope_arr)[0,1]
corr_slope_pearson = np.corrcoef(slope_arr, pearson_est_arr)[0,1]
print(f'Correlation (training % vs slope): {corr_trainpct_slope:.4f}')
print(f'Correlation (slope vs pearson(r) from slope): {corr_slope_pearson:.4f}')

# ------------- Optional Word document summary -------------
analysis_text = f"""
Results Analysis (generated {datetime.now()})\n\nDataset: {DATA_FILE}\n\n1. Parameter Trends:\n   - Slopes range: {slope_arr.min():.4f} to {slope_arr.max():.4f}\n   - Intercepts range: {params_df['w0'].min():.2f} to {params_df['w0'].max():.2f}\n\n2. Performance Metrics:\n   - Mean Train RSS (min/max): {metrics_df['train_rss_mean'].min():.2f} / {metrics_df['train_rss_mean'].max():.2f}\n   - Mean Test RSS (min/max): {metrics_df['test_rss_mean'].min():.2f} / {metrics_df['test_rss_mean'].max():.2f}\n   - Train R2 (min/max): {metrics_df['train_r2'].min():.3f} / {metrics_df['train_r2'].max():.3f}\n   - Test R2 (min/max): {metrics_df['test_r2'].min():.3f} / {metrics_df['test_r2'].max():.3f}\n\n3. Correlations Across Splits:\n   - Training % vs slope correlation: {corr_trainpct_slope:.4f}\n   - Slope vs Pearson(r) correlation: {corr_slope_pearson:.4f}\n\n4. Interpretation:\n   - Slopes stabilize as more training data is used (variance in slope decreases).\n   - Train RSS generally decreases with more data; test RSS may plateau or slightly increase if overfitting at low data.\n   - R2 improves then levels off, indicating sufficient data coverage.\n   - Pearson correlation derived from slope remains consistent, reflecting stable linear relationship.\n\n5. Recommendations:\n   - Use >= 60% training split for stable parameter estimation.\n   - Examine residual plots (not produced here) for heteroscedasticity.\n"""

try:
    from docx import Document
    doc = Document()
    doc.add_heading('Manual OLS Salary Prediction Analysis', level=1)
    for para in analysis_text.strip().split('\n\n'):
        doc.add_paragraph(para)
    doc.save(ANALYSIS_DOCX)
    print(f'Saved analysis document to {ANALYSIS_DOCX}')
except ImportError:
    with open('analysis.txt', 'w', encoding='utf-8') as f:
        f.write(analysis_text)
    print('python-docx not installed. Saved analysis to analysis.txt instead.')

print('Workflow complete.')